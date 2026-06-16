import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PURPLE = RGBColor(0x7B, 0x2D, 0x8B)
TEAL   = RGBColor(0x00, 0xA8, 0x9D)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1A, 0x1A, 0x1A)


def set_cell_bg(cell, hex_color):
    # Word doesn't support cell background colors through the standard python-docx API,
    # so we drop down to raw XML to set the shading element.
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_para(cell, text, bold=False, color=DARK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()
    para.alignment = align
    run = para.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return para


def stars_text(rating):
    if not rating:
        return "N/A"
    try:
        return str(int(rating))
    except Exception:
        return str(rating)


def add_section_row(table, label_text, label_bg="7B2D8B"):
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_bg(cell, label_bg)
    add_para(cell, label_text, bold=True, color=WHITE, size=9,
             align=WD_ALIGN_PARAGRAPH.CENTER)


def add_data_row(table, label, value, row_idx=0):
    row = table.add_row()
    lc = row.cells[0]
    vc = row.cells[1]
    set_cell_bg(lc, "EFEFEF")
    set_cell_bg(vc, "FFFFFF" if row_idx % 2 == 0 else "FAFAFA")
    set_cell_borders(lc)
    set_cell_borders(vc)
    add_para(lc, label, bold=True, color=DARK, size=9)
    add_para(vc, value, color=DARK, size=9)


def generate_docx(data: dict, buf):
    doc = Document()

    section = doc.sections[0]
    section.left_margin  = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin   = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Logo header - falls back to text if logo.png is missing
    logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(2.8))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("INFINITE")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = PURPLE
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Managed by MEDELITE")
        run2.font.size = Pt(11)
        run2.font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("FACILITY ASSESSMENT SNAPSHOT")
    run3.bold = True
    run3.font.size = Pt(13)
    run3.font.color.rgb = PURPLE

    state = data.get("state", "")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(state)
    run4.bold = True
    run4.font.size = Pt(11)

    ccn = data.get("ccn", "")
    medicare_url = f"https://www.medicare.gov/care-compare/details/nursing-home/{ccn}"
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run(f"Medicare Care Compare: {medicare_url}")
    run5.font.size = Pt(8)
    run5.font.color.rgb = TEAL

    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(2.6)
    table.columns[1].width = Inches(4.6)

    display_name = data.get("name_override") or data.get("name") or ""
    address_full = ", ".join(filter(None, [
        data.get("address", ""),
        data.get("city", ""),
        data.get("state", ""),
        data.get("zip", ""),
    ]))

    basic_rows = [
        ("Name of Facility",                           display_name),
        ("Location",                                   address_full),
        ("EMR",                                        data.get("emr", "")),
        ("Census Capacity",                            str(data.get("certified_beds", ""))),
        ("Current Census",                             str(data.get("current_census", ""))),
        ("Type of Patient",                            data.get("patient_type", "")),
        ("Previous Coverage from Medelite",            data.get("prev_coverage", "")),
        ("Previous Provider Performance from Medelite", data.get("prev_performance", "")),
        ("Medical Coverage",                           data.get("medical_coverage", "")),
    ]
    for i, (lbl, v) in enumerate(basic_rows):
        add_data_row(table, lbl, v, row_idx=i)

    add_section_row(table, "STAR RATINGS - CMS Five-Star Quality Rating System", "7B2D8B")
    star_rows = [
        ("Overall Star Rating",        stars_text(data.get("overall_rating"))),
        ("Health Inspection",          stars_text(data.get("health_inspection_rating"))),
        ("Staffing",                   stars_text(data.get("staffing_rating"))),
        ("Quality of Resident Care",   stars_text(data.get("qm_rating"))),
    ]
    for i, (lbl, v) in enumerate(star_rows):
        add_data_row(table, lbl, v, row_idx=i)

    add_section_row(table, "HOSPITALIZATION & ED METRICS", "00A89D")
    hosp_rows = [
        ("Short Term Hospitalization",          data.get("str_hosp", "")),
        ("STR National Avg. for Hospitalization", data.get("str_hosp_nat", "")),
        ("STR State Avg. for Hospitalization",  data.get("str_hosp_state", "")),
        ("STR ED Visit",                        data.get("str_ed", "")),
        ("STR ED Visits National Avg.",         data.get("str_ed_nat", "")),
        ("STR ED Visits State Avg.",            data.get("str_ed_state", "")),
        ("LT Hospitalization",                  data.get("lt_hosp", "")),
        ("LT National Avg. for Hospitalization", data.get("lt_hosp_nat", "")),
        ("LT State Avg. for Hospitalization",   data.get("lt_hosp_state", "")),
        ("ED Visit (LT per 1000)",              data.get("lt_ed", "")),
        ("LT ED Visits National Avg.",          data.get("lt_ed_nat", "")),
        ("LT ED Visits State Avg.",             data.get("lt_ed_state", "")),
    ]
    for i, (lbl, v) in enumerate(hosp_rows):
        add_data_row(table, lbl, v, row_idx=i)

    doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(
        f"Data sourced from CMS Provider Data Catalog. "
        f"Report generated by INFINITE, Managed by MEDELITE. CCN: {ccn}"
    )
    rf.font.size = Pt(7)
    rf.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(buf)
